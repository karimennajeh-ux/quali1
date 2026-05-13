from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
REPORT_DIR = ROOT / "database" / "word_documents" / "reports"
ASSET_DIR = REPORT_DIR / "Mode_operatoire_xampp_assets"
OUT = REPORT_DIR / "Mode_operatoire_utilisation_XAMPP_sauvegarde_automatique.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Segoe UI"
        run.font.color.rgb = RGBColor(31, 78, 121)
    return p


def add_note(doc, title, body, fill="EAF3FF"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Segoe UI"
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(31, 78, 121)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.name = "Segoe UI"
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_picture(doc, image_name, caption):
    path = ASSET_DIR / image_name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Cm(15.5))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.style = "Caption"


def add_checklist(doc, items):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(1.2)
    table.columns[1].width = Cm(14.5)
    set_cell_text(table.rows[0].cells[0], "OK", True, "FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Controle a effectuer", True, "FFFFFF")
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "1F4E79")
    for item in items:
        row = table.add_row()
        set_cell_text(row.cells[0], "[ ]", False)
        set_cell_text(row.cells[1], item, False)
    doc.add_paragraph()


def add_data_table(doc):
    rows = [
        ("Comptes pilotes", "SQLite table pilots", "Creation, modification, suppression, mot de passe hashe."),
        ("Utilisateurs", "SQLite table users", "Comptes agents, profils, permissions et acces par module."),
        ("Donnees metier", "SQLite table pilot_app_state", "Tableau de bord, documents, equipements, NC, audits, clients, planning, statistiques, etc."),
        ("Documentation centrale", "SQLite + QUALI_DATA_SERVER", "Metadonnees en base, fichiers stockes dans le depot documentaire local."),
        ("Cache navigateur", "localStorage", "Copie de secours locale; la source serveur est resynchronisee a la connexion."),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Donnees", "Emplacement serveur local", "Role"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, True, "FFFFFF")
        set_cell_shading(table.rows[0].cells[i], "1F4E79")
    for data, place, role in rows:
        row = table.add_row()
        set_cell_text(row.cells[0], data)
        set_cell_text(row.cells[1], place)
        set_cell_text(row.cells[2], role)
    doc.add_paragraph()


def build():
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Segoe UI"
    styles["Normal"].font.size = Pt(10.5)
    styles["Caption"].font.name = "Segoe UI"
    styles["Caption"].font.size = Pt(9)
    styles["Caption"].font.italic = True

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Mode operatoire d'utilisation")
    r.bold = True
    r.font.name = "Segoe UI"
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(31, 78, 121)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("QUALI by ENNAJEH - XAMPP local + sauvegarde automatique")
    sr.font.name = "Segoe UI"
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(85, 85, 85)

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta_rows = [
        ("Application", "QUALI by ENNAJEH v2"),
        ("URL utilisateur", "http://localhost/QUALI/"),
        ("API locale", "http://localhost:3000/api"),
        ("Date", "13/05/2026"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        set_cell_text(meta.rows[i].cells[0], k, True)
        set_cell_shading(meta.rows[i].cells[0], "EAF3FF")
        set_cell_text(meta.rows[i].cells[1], v)
    doc.add_paragraph()

    add_note(
        doc,
        "Objectif",
        "Ce document explique comment lancer l'application avec XAMPP comme serveur local, "
        "comment se connecter, et comment verifier que les donnees saisies sont enregistrees automatiquement "
        "dans le backend local Node/SQLite.",
    )

    add_heading(doc, "1. Architecture locale", 1)
    p = doc.add_paragraph()
    p.add_run("Principe retenu : ").bold = True
    p.add_run(
        "Apache/XAMPP sert l'interface a l'adresse http://localhost/QUALI/. "
        "Le backend Node reste actif sur le port 3000 et centralise les donnees dans SQLite. "
        "L'interface appelle automatiquement http://localhost:3000/api pour charger et sauvegarder l'etat du pilote."
    )
    add_data_table(doc)

    add_heading(doc, "2. Lancement", 1)
    add_checklist(
        doc,
        [
            "Demarrer Apache dans XAMPP ou executer scripts/launch_quali_local.ps1.",
            "Verifier que le backend Node est lance sur le port 3000.",
            "Ouvrir http://localhost/QUALI/ dans le navigateur.",
            "Verifier que la page affiche QUALI by ENNAJEH et non le JSON /api/health.",
        ],
    )
    add_picture(doc, "01_accueil_xampp.png", "Figure 1 - Application ouverte via XAMPP sur http://localhost/QUALI/.")

    add_heading(doc, "3. Connexion", 1)
    doc.add_paragraph(
        "Cliquer sur Continuer, renseigner l'adresse e-mail du compte pilote et le mot de passe, "
        "puis cliquer sur Se connecter. Les mots de passe ne sont pas renvoyes par l'API et sont stockes sous forme de hash cote serveur."
    )
    add_picture(doc, "02_connexion_xampp.png", "Figure 2 - Ecran de connexion servi par XAMPP.")

    add_heading(doc, "4. Sauvegarde automatique", 1)
    doc.add_paragraph(
        "Chaque action met d'abord a jour l'affichage local, puis declenche une sauvegarde serveur environ une demi-seconde plus tard. "
        "Les actions concernees incluent notamment les documents, equipements, non-conformites, audits, clients, prestataires, planning, "
        "statistiques, diagnostic, SWOT, satisfaction, reclamations, personnel, messages et parametres organisme."
    )
    add_note(
        doc,
        "Bon reflexe",
        "Apres une creation ou modification, attendre le message 'Sauvegarde serveur local effectuee'. "
        "Si le serveur Node est arrete, l'application garde une copie locale et affiche que le serveur local est indisponible.",
        "FFF4DF",
    )
    add_picture(doc, "03_tableau_bord.png", "Figure 3 - Tableau de bord apres connexion; les modules utilisent la sauvegarde automatique.")

    add_heading(doc, "5. Verification technique", 1)
    add_checklist(
        doc,
        [
            "http://localhost/QUALI/ doit retourner la page de l'application.",
            "http://localhost:3000/api/health doit retourner ok=true.",
            "Dans l'API health, le champ database doit pointer vers database/qualilab.sqlite.",
            "Le dossier database/QUALI_DATA_SERVER contient le depot documentaire central.",
            "Ne pas ouvrir /api/health comme application: c'est seulement une page de controle technique.",
        ],
    )

    add_heading(doc, "6. Depannage rapide", 1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Symptome", True, "FFFFFF")
    set_cell_text(table.rows[0].cells[1], "Action", True, "FFFFFF")
    set_cell_shading(table.rows[0].cells[0], "1F4E79")
    set_cell_shading(table.rows[0].cells[1], "1F4E79")
    tips = [
        ("Page inaccessible", "Demarrer Apache/XAMPP, puis verifier le port 80."),
        ("Erreur serveur local indisponible", "Demarrer Node avec npm start ou scripts/launch_quali_local.ps1."),
        ("L'URL affiche du JSON", "Revenir a http://localhost/QUALI/ au lieu de /api/health."),
        ("Les donnees ne se synchronisent pas", "Verifier /api/health et attendre le message de sauvegarde serveur."),
    ]
    for a, b in tips:
        row = table.add_row()
        set_cell_text(row.cells[0], a)
        set_cell_text(row.cells[1], b)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
